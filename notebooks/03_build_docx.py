"""
03_build_docx.py
=================
Reproducible builder for the journal-ready DOCX export of P1.

Converts paper/P1_manuscript.md (plus the committed PNG figures in figures/)
into paper/P1_manuscript_JRFM.docx, styled to MDPI / Journal of Risk and
Financial Management (JRFM) conventions. This mirrors the P2 deliverable
P2_manuscript_JRFM.docx and keeps the repository as the single source of truth:
the DOCX is a build artifact, never hand-edited.

Usage (from repository root):
    python3 notebooks/03_build_docx.py

Primary path uses Pandoc (preferred, best Markdown fidelity):
    pandoc paper/P1_manuscript.md -o paper/P1_manuscript_JRFM.docx \\
        --resource-path=figures --reference-doc=<optional template>
If Pandoc is unavailable, a pure-Python fallback using python-docx renders the
manuscript and embeds the figures at 300 DPI-equivalent width.

Dependencies (see environments/ / requirements):
    - pandoc (system binary)            # primary
    - python-docx, markdown, Pillow     # fallback

No network access and no fabricated data: every byte of content comes from the
committed manuscript and the committed figures.
"""
from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

# --- Repository-relative paths (script lives in notebooks/) ------------------
REPO_ROOT = Path(__file__).resolve().parents[1]
PAPER_DIR = REPO_ROOT / "paper"
FIG_DIR = REPO_ROOT / "figures"
MD_PATH = PAPER_DIR / "P1_manuscript.md"
DOCX_PATH = PAPER_DIR / "P1_manuscript_JRFM.docx"

# Expected figures, matching the ![](../figures/...) embeds in the manuscript.
FIGURES = [
    "fig1_wacc_firr_gap.png",
    "fig2_ridership_forecast.png",
    "fig3_revenue_forecast.png",
    "fig4_firr_sensitivity_tornado.png",
]


def _check_inputs() -> None:
    """Fail fast if any source-of-truth input is missing."""
    if not MD_PATH.exists():
        sys.exit(f"ERROR: manuscript not found: {MD_PATH}")
    missing = [f for f in FIGURES if not (FIG_DIR / f).exists()]
    if missing:
        sys.exit(f"ERROR: missing committed figures: {missing}")
    print(f"OK  manuscript: {MD_PATH.relative_to(REPO_ROOT)}")
    print(f"OK  figures:    {len(FIGURES)} present in {FIG_DIR.relative_to(REPO_ROOT)}")


def build_with_pandoc() -> bool:
    """Primary path: convert via Pandoc. Returns True on success."""
    if shutil.which("pandoc") is None:
        print("INFO  pandoc not found on PATH; using python-docx fallback.")
        return False
    cmd = [
        "pandoc",
        str(MD_PATH),
        "-o", str(DOCX_PATH),
        "--resource-path", str(FIG_DIR),
        "--from", "gfm",
        "--standalone",
    ]
    print("RUN  " + " ".join(cmd))
    result = subprocess.run(cmd, cwd=str(REPO_ROOT))
    return result.returncode == 0


def build_with_python_docx() -> bool:
    """Fallback path: render with python-docx (MDPI/JRFM-ish styling)."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        sys.exit("ERROR: install pandoc OR `pip install python-docx Pillow`.")

    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()

    # Base style: 10pt Palatino Linotype-like body, JRFM convention.
    normal = doc.styles["Normal"]
    normal.font.name = "Palatino Linotype"
    normal.font.size = Pt(10)

    fig_embed = re.compile(r"!\[[^\]]*\]\(\.\./figures/([^)]+)\)")
    bold_inline = re.compile(r"\*\*(.+?)\*\*")

    def add_runs(p, line):
        """Render **bold** spans inside an otherwise plain paragraph."""
        pos = 0
        for m in bold_inline.finditer(line):
            if m.start() > pos:
                p.add_run(line[pos:m.start()])
            r = p.add_run(m.group(1))
            r.bold = True
            pos = m.end()
        if pos < len(line):
            p.add_run(line[pos:])

    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        m = fig_embed.search(line)
        if m:
            fig_file = m.group(1).split("\\")[0]
            fig_path = FIG_DIR / fig_file
            if fig_path.exists():
                doc.add_picture(str(fig_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.strip() == "---":
            continue
        else:
            add_runs(doc.add_paragraph(), line)

    doc.save(str(DOCX_PATH))
    return True


def main() -> None:
    print("=== P1 JRFM DOCX builder ===")
    _check_inputs()
    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    ok = build_with_pandoc()
    if not ok:
        ok = build_with_python_docx()
    if ok and DOCX_PATH.exists():
        size_kb = DOCX_PATH.stat().st_size / 1024
        print(f"DONE  wrote {DOCX_PATH.relative_to(REPO_ROOT)} ({size_kb:.1f} KB)")
    else:
        sys.exit("ERROR: DOCX build failed.")


if __name__ == "__main__":
    main()
